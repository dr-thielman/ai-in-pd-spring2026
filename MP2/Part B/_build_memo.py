"""Build MP2_PartB_Research_Directive.docx — Jordan Chen's in-world memo.

Run from repo root:
    python "MP2/Part B/_build_memo.py"
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "MP2_PartB_Research_Directive.docx")


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_body(doc, text, spacing_after=Pt(6)):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = spacing_after
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_bold_para(doc, bold_text, rest=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if rest:
        run2 = p.add_run(rest)
        run2.font.size = Pt(11)
        run2.font.name = "Calibri"
    return p


# ── Build document ───────────────────────────────────────────────────────

doc = Document()

# -- Page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# -- Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.line_spacing = Pt(15)

# ═══════════════════════════════════════════════════════════════════════════
# HEADER
# ═════════════════════════���═════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MEMORANDUM")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Horizontal rule
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
pPr = p._element.get_or_add_pPr()
pBdr = pPr.makeelement(qn("w:pBdr"), {})
bottom = pBdr.makeelement(qn("w:bottom"), {
    qn("w:val"): "single",
    qn("w:sz"): "6",
    qn("w:space"): "1",
    qn("w:color"): "1A1A2E",
})
pBdr.append(bottom)
pPr.append(pBdr)

# -- Memo metadata table (no borders)
meta = [
    ("To:", "Engineering Team"),
    ("From:", "Jordan Chen, Engineering Manager, Gripper Systems Division"),
    ("Date:", "April 14, 2026"),
    ("Re:", "MiniClaw Project \u2014 Research Phase Before Detailed Design"),
]
table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
# Remove borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
borders = tblPr.makeelement(qn("w:tblBorders"), {})
for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
    border = borders.makeelement(qn(f"w:{edge}"), {
        qn("w:val"): "none",
        qn("w:sz"): "0",
        qn("w:space"): "0",
        qn("w:color"): "auto",
    })
    borders.append(border)
tblPr.append(borders)

for label, value in meta:
    row = table.add_row()
    # Label
    cell_l = row.cells[0]
    p = cell_l.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    cell_l.width = Inches(0.8)
    # Value
    cell_r = row.cells[1]
    p = cell_r.paragraphs[0]
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    cell_r.width = Inches(5.7)

# Another horizontal rule
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(8)
pPr = p._element.get_or_add_pPr()
pBdr = pPr.makeelement(qn("w:pBdr"), {})
bottom = pBdr.makeelement(qn("w:bottom"), {
    qn("w:val"): "single",
    qn("w:sz"): "6",
    qn("w:space"): "1",
    qn("w:color"): "1A1A2E",
})
pBdr.append(bottom)
pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════
# BODY
# ═════════════════���═════════════════════════════════════════════════════════

# Opening
add_body(doc,
    "I\u2019ve reviewed the gear train designs from the initial MiniClaw assignment, "
    "and I want to start with the good news: this team showed real initiative. The "
    "AI-assisted approach generated concepts fast, and several of you came up with "
    "creative gear train configurations I wouldn\u2019t have thought of immediately. "
    "There\u2019s genuine engineering talent in these submissions."
)

add_body(doc,
    "Now the harder conversation. Looking at the designs more carefully, I\u2019m seeing "
    "the same pattern I see in early-career engineers who jump to solutions before "
    "doing their homework. Too many of these designs were built on generic data "
    "instead of ACME\u2019s actual capabilities, and that means rework."
)

# What concerned Jordan
add_bold_para(doc, "What I\u2019m seeing")

add_body(doc,
    "Several designs used PLA material properties pulled straight from the internet \u2014 "
    "datasheets from random filament brands, generic FDM specifications, material "
    "values that don\u2019t match what comes off our printers. We run Prusa MK4S machines "
    "with 0.4 mm nozzles, printing FilaTech PolyPro PLA+. That\u2019s a specific material "
    "with specific properties that we\u2019ve actually tested in our lab. The interlayer "
    "adhesion number alone is significantly different from what you\u2019ll find on a "
    "generic PLA datasheet, and for gear teeth, that\u2019s the number that matters."
)

add_body(doc,
    "Some designs specified tolerances our print shop can\u2019t hold, or used gear "
    "parameters that our internal guidelines specifically warn against. A few of you "
    "designed every component from scratch when ACME already has a component library "
    "from the WidgetBot and GripperBot projects \u2014 brass pin stock, thumb wheel blanks, "
    "silicone grip pads, all sitting on the shelf. Nobody asked about our vendor "
    "relationship with FilaTech or what the fab team\u2019s actual tolerance capabilities are."
)

add_body(doc,
    "In industry, jumping to design before doing your homework costs real money. We "
    "caught it early this time. Let\u2019s do the research phase right."
)

# What Jordan wants
add_bold_para(doc, "What I need from you")

add_body(doc,
    "I\u2019m sharing access to the ACME internal knowledge base \u2014 about 15 documents "
    "covering our manufacturing capabilities, material test data, previous product "
    "history, engineering design standards, vendor communications, and project status. "
    "This is the information your designs should have been built on. Use it."
)

add_body(doc, "In two weeks, I need the following from each of you:")

p = doc.add_paragraph(style="List Bullet")
run = p.add_run(
    "A research package that demonstrates you\u2019ve actually reviewed the company "
    "knowledge base. I want to see evidence that you understand ACME\u2019s manufacturing "
    "capabilities, material properties, and lessons learned from previous projects."
)
run.font.size = Pt(11)

p = doc.add_paragraph(style="List Bullet")
run = p.add_run(
    "An honest assessment of what you know versus what you still need to find out. "
    "I value intellectual honesty over false confidence. If you don\u2019t know something, "
    "say so and explain how you\u2019d find the answer."
)
run.font.size = Pt(11)

p = doc.add_paragraph(style="List Bullet")
run = p.add_run(
    "A working set of product requirements grounded in ACME\u2019s actual capabilities. "
    "The requirements from MP1 were a starting point, but now they should reflect our "
    "real print tolerances, our tested material properties, and our production constraints."
)
run.font.size = Pt(11)

p = doc.add_paragraph(style="List Bullet")
run = p.add_run(
    "A preliminary design direction with rationale. Not a final design \u2014 a well-reasoned "
    "concept that accounts for what you\u2019ve learned from the ACME knowledge base."
)
run.font.size = Pt(11)

# Specific technical context
add_bold_para(doc, "Key context you should know")

add_body(doc,
    "Our print shop runs 12 Prusa MK4S printers with 0.4 mm nozzles as standard. "
    "Designs need to respect their actual capabilities, not generic FDM specs from "
    "the internet. Maria Santos and the fab team have documented everything \u2014 "
    "achievable tolerances, press-fit guidelines, pin minimums, print orientation rules. "
    "It\u2019s all in the knowledge base."
)

add_body(doc,
    "We\u2019ve printed gears before. The WidgetBot 2.0 project used PLA spur gears, and "
    "Elena Vasquez ran a full fatigue test program on them. There\u2019s test data in the "
    "company files, including failure analysis at various torque levels. The results will "
    "change how you think about safety factors for printed gears. Read the test report."
)

add_body(doc,
    "The RobotExpo deadline is firm: June 10\u201312. The fab team needs final design files "
    "by May 15 to produce 500+ units on our printer fleet. That gives us about a month "
    "between your research phase and the detailed design that follows. The research "
    "phase needs to happen now \u2014 there is no slack in this schedule."
)

# Closing
add_bold_para(doc, "Looking ahead")

add_body(doc,
    "I\u2019ll review the research packages when they come in and then greenlight the "
    "detailed design phase. Think of this as doing the homework so that the next phase "
    "goes smoothly. It\u2019s a lot cheaper to spend two weeks researching than two weeks "
    "reworking a design that was built on bad assumptions."
)

add_body(doc,
    "The MiniClaw is a strategic project with visibility to our VP of Engineering. "
    "Let\u2019s get the foundation right."
)

# Signature
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Jordan Chen")
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph("Engineering Manager, Gripper Systems Division")
for run in p.runs:
    run.font.size = Pt(11)

p = doc.add_paragraph("ACME Robotics")
for run in p.runs:
    run.font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══��════════════════��════════════════════════════��═════════════════════════

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
