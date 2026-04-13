"""Build MP2_Context_Management.docx — the assignment document for MP2.

Run from repo root:
    python "MP2/Part B/_build_assignment.py"
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "MP2_Context_Management.docx")


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_table_row(table, cells_data, bold=False, shading=None):
    """Add a row to a table with optional formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        if bold:
            run.bold = True
        if shading:
            set_cell_shading(cell, shading)
    return row


def add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_body(doc, text):
    """Add a body paragraph."""
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_bold_body(doc, bold_text, rest_text=""):
    """Add a paragraph starting with bold text."""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(11)
    if rest_text:
        run2 = p.add_run(rest_text)
        run2.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


# ── Build document ───────��───────────────────────────────────────────────

doc = Document()

# -- Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# -- Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══���═════════════════���══════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MINI-PROJECT 2")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The Research Phase")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Context Management for AI-Augmented Engineering")
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

doc.add_paragraph()  # spacer

# -- Metadata table
meta = [
    ("Type", "Individual"),
    ("Points", "100 points (Part A: 50 pts | Part B: 50 pts)"),
    ("Due Date", "Monday, April 27, 2026 at 11:59 PM (both Part A and Part B)"),
    ("Primary Pillar", "Context Management \u2605\u2605\u2605"),
    ("Tools", "GitHub Codespaces, GitHub Copilot Pro, ChromaDB, GitHub Models API, Frontier AI"),
]
table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for label, value in meta:
    row = table.add_row()
    # Label cell
    cell_l = row.cells[0]
    p = cell_l.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(cell_l, "E8EDF5")
    # Value cell
    cell_r = row.cells[1]
    p = cell_r.paragraphs[0]
    run = p.add_run(value)
    run.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.0)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ══��═════════════════��═════════════════════════════��════════════════════════

add_heading(doc, "Overview", level=1)

add_body(doc,
    "This project has two parts that teach you how to manage the information "
    "available to AI systems \u2014 and why it matters for engineering work."
)

p = doc.add_paragraph()
run = p.add_run("Part A")
run.bold = True
run.font.size = Pt(11)
run = p.add_run(
    " is a hands-on notebook where you build a complete RAG (Retrieval-Augmented "
    "Generation) pipeline: tokenizing documents, embedding them, storing them in "
    "a vector database, and querying them with an LLM. You\u2019ll end by querying "
    "the actual \u201CAttention Is All You Need\u201D paper using the technique it describes."
)
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("Part B")
run.bold = True
run.font.size = Pt(11)
run = p.add_run(
    " is the research phase of the MiniClaw project. Jordan Chen has reviewed your "
    "MP1 submissions and wants deeper research before moving to detailed design. "
    "Your job: build a context package of ACME\u2019s internal documents, demonstrate "
    "that context management improves AI engineering answers, and produce a working "
    "PRD that will guide MP3\u2019s detailed design."
)
run.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════════════════
# PART A
# ══════════════════════���════════════════════════════════════════════════════

add_heading(doc, "Part A: The Information Problem", level=1)

p = doc.add_paragraph()
run = p.add_run("Due: Monday, April 27 at 11:59 PM  |  50 points  |  ~60\u201390 minutes")
run.italic = True
run.font.size = Pt(11)

add_body(doc,
    "Open the provided Jupyter notebook in your GitHub Codespace and work through "
    "six sections. Build a RAG pipeline section by section: tokens, embeddings, "
    "chunking, vector storage, retrieval, and generation. Each section has a point "
    "value and specific output to produce."
)

add_bold_body(doc, "Submit: ",
    "Completed notebook pushed to your course repository (all cells run, outputs visible)."
)

add_body(doc,
    "The RAG pipeline you build in Part A is one approach for managing the ACME "
    "context. You may use that pipeline, adapt it, or use any approach you choose "
    "for Part B."
)

# ══════════════════════════════════════════���════════════════════════════════
# PART B
# ════════════════════════════���═══════════════════════════���══════════════════

add_heading(doc, "Part B: Context Management for the MiniClaw", level=1)

p = doc.add_paragraph()
run = p.add_run("Due: Monday, April 27 at 11:59 PM  |  50 points")
run.italic = True
run.font.size = Pt(11)

add_body(doc,
    "Read the attached memo from Jordan Chen. Then, using the ACME document corpus "
    "provided and the RAG skills from Part A, build the context management package "
    "for the MiniClaw project."
)

add_body(doc,
    "Open the provided starter notebook (MP2_PartB_Context_Management.ipynb) in your "
    "GitHub Codespace. The notebook has section headers and the corpus loading code "
    "pre-written. You fill in everything else \u2014 use markdown cells for written "
    "deliverables and code cells for your RAG pipeline work."
)

add_bold_body(doc, "Submit: ",
    "Completed notebook pushed to your course repository (all cells run, outputs visible). "
    "Same process as Part A \u2014 commit, push, and submit your repo URL on Canvas."
)

# -- Deliverables
add_heading(doc, "Deliverables", level=2)

# Deliverable 1
add_bold_body(doc, "1. Project Intake Document (10 pts)")
add_body(doc,
    "Apply the PCS intake framework (from Session 5) to the MiniClaw project. This "
    "is the document that would kick off the project in a real engineering firm. "
    "Include: vision (from ACME\u2019s perspective), mission (your engineering scope), "
    "what\u2019s in/out of scope, key assumptions, responsibilities, and a Phase Zero "
    "recommendation (or rationale for proceeding directly to development)."
)
add_body(doc,
    "Note: The MiniClaw design brief from MP1 provides the starting point, but your "
    "intake should reflect deeper understanding from the ACME corpus \u2014 company "
    "context, manufacturing capabilities, business goals for RobotExpo."
)

# Deliverable 2
add_bold_body(doc, "2. Information Strategy (8 pts)")
add_body(doc,
    "Answer: what information do you need to design the MiniClaw well? "
    "Categorize into three buckets:"
)
add_bullet(doc,
    "What the AI already knows well (general engineering knowledge, standard "
    "formulas, common materials)"
)
add_bullet(doc,
    "What the AI knows partially (PLA printing parameters, gear design "
    "guidelines \u2014 correct in general, may be wrong in specifics)"
)
add_bullet(doc,
    "What requires ACME-specific context (internal test data, manufacturing "
    "tolerances, vendor relationships, business constraints)"
)
add_body(doc,
    "For each bucket, give 2\u20133 specific examples relevant to the MiniClaw. "
    "This demonstrates you understand WHY context management matters, not just "
    "HOW to do it."
)

# Deliverable 3
add_bold_body(doc, "3. Context Package + Before/After Demo (15 pts)")
add_body(doc,
    "This is the core deliverable. Build a context package by loading the ACME "
    "document corpus into a RAG pipeline (using the approach from Part A or any "
    "approach you choose)."
)
add_body(doc, "Then demonstrate the value of context management:")
add_bullet(doc, "Choose 3 specific MiniClaw engineering questions where ACME context would matter")
add_bullet(doc, "For each question, show:")

p = doc.add_paragraph(style="List Bullet 2")
run = p.add_run("The AI\u2019s answer WITHOUT the ACME context (just general knowledge)")
run.font.size = Pt(11)

p = doc.add_paragraph(style="List Bullet 2")
run = p.add_run("The AI\u2019s answer WITH the ACME context (RAG-augmented)")
run.font.size = Pt(11)

p = doc.add_paragraph(style="List Bullet 2")
run = p.add_run(
    "Your assessment: what changed? Is the context-augmented answer better? "
    "What did it get right that the baseline missed?"
)
run.font.size = Pt(11)

add_body(doc,
    "The questions should be meaningful engineering questions, not trivia. "
    "Good example: \u201CWhat PLA printing parameters should I use for the MiniClaw "
    "gears?\u201D Bad example: \u201CWhat is ACME\u2019s address?\u201D"
)

# Deliverable 4
add_bold_body(doc, "4. Research Synthesis (7 pts)")
add_body(doc,
    "Use AI (with and without your context package) to research 2\u20133 technical "
    "questions relevant to your MiniClaw gear train design. For each:"
)
add_bullet(doc, "The question you asked")
add_bullet(doc, "The tool(s) you used and the context you provided")
add_bullet(doc, "The answer you received")
add_bullet(doc, "Your engineering evaluation: do you trust this? What would you verify?")
add_body(doc,
    "These should advance your actual MiniClaw design \u2014 this isn\u2019t make-work, "
    "it\u2019s preparation for MP3."
)

# Deliverable 5
add_bold_body(doc, "5. Preliminary Design Concept / Working PRD (10 pts)")
add_body(doc,
    "Based on your research and context work, produce a working PRD for the "
    "MiniClaw. This is not a final design \u2014 it\u2019s a requirements document that "
    "will guide MP3\u2019s detailed design phase. Include:"
)
add_bullet(doc,
    "8\u201312 product requirements (specific, measurable, testable \u2014 as practiced "
    "in Session 6)"
)
add_bullet(doc,
    "For each requirement, note whether it came from the design brief, your "
    "research, or the ACME context"
)
add_bullet(doc, "2\u20133 design directions you\u2019re considering, with brief rationale")
add_bullet(doc, "Key open questions that MP3 will need to resolve")
add_body(doc,
    "Your PRD will guide your detailed design work in MP3. Write it as if your "
    "future self is the audience."
)

# ══════════════════════════════════���════════════════════════════════════════
# GRADING RUBRIC
# ════════════════════════════════���══════════════════════════════════════════

add_heading(doc, "Grading Rubric", level=1)

# -- Part A rubric
add_heading(doc, "Part A: The Information Problem (50 points)", level=2)

tbl_a = doc.add_table(rows=1, cols=3)
tbl_a.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header
for i, h in enumerate(["Component", "Points", "What We Are Looking For"]):
    cell = tbl_a.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, "2B3A67")

rows_a = [
    ("Completion & Experimentation", "30",
     "All cells run. Code cells show working implementations. Outputs match expected formats."),
    ("Reflections", "20",
     "Three thoughtful reflections demonstrating observation and connection to course concepts."),
]
for cells_data in rows_a:
    add_table_row(tbl_a, cells_data)

# Column widths
for row in tbl_a.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(0.7)
    row.cells[2].width = Inches(3.8)

doc.add_paragraph()  # spacer

# -- Part B rubric
add_heading(doc, "Part B: Context Management for the MiniClaw (50 points)", level=2)

add_body(doc,
    "Every mini-project is evaluated across all five pillars. The weight shifts "
    "based on the project\u2019s primary emphasis."
)

tbl_b = doc.add_table(rows=1, cols=3)
tbl_b.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header
for i, h in enumerate(["Pillar", "Points", "What We Are Looking For"]):
    cell = tbl_b.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, "2B3A67")

rows_b = [
    ("Context Management \u2605\u2605\u2605", "20",
     "Before/after demo shows meaningful improvement. Information strategy demonstrates "
     "understanding of what AI knows vs. doesn\u2019t. Context package is well-constructed. "
     "ACME documents used effectively."),
    ("Goal & Direction \u2605\u2605", "10",
     "Intake document is thorough. PRD requirements are specific, measurable, testable. "
     "Research questions are well-chosen and advance the design."),
    ("Evaluation & Trust \u2605\u2605", "10",
     "Before/after assessments are specific and honest. Research synthesis includes trust "
     "evaluation. Student can articulate what needs verification."),
    ("Centaur Engineering \u2605", "5",
     "Evidence of genuine back-and-forth with AI. Student directed the AI, didn\u2019t just "
     "accept outputs."),
    ("Tools & Integration \u2605", "5",
     "RAG pipeline used appropriately. Right tool for each sub-task."),
    ("Total", "50", ""),
]
for cells_data in rows_b:
    row = add_table_row(tbl_b, cells_data)
    if cells_data[0] == "Total":
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

# Column widths
for row in tbl_b.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(0.7)
    row.cells[2].width = Inches(3.8)

# ══════════════════════════════════════════════════════════════��════════════
# REFERENCE: KEY GEAR EQUATIONS
# ═══════════════���═══════════════════════════════════════════════════════════

add_heading(doc, "Reference: Key Gear Equations", level=1)

add_body(doc,
    "You are expected to use AI to help with these, but you should recognize them "
    "and understand what they mean."
)

tbl_ref = doc.add_table(rows=1, cols=2)
tbl_ref.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Relationship", "Formula / Description"]):
    cell = tbl_ref.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, "2B3A67")

refs = [
    ("Pitch Diameter", "d = m \u00d7 z   (module \u00d7 tooth count)"),
    ("Center Distance", "a = m \u00d7 (z\u2081 + z\u2082) / 2"),
    ("Gear Ratio", "i = z\u2082 / z\u2081 = T\u2082 / T\u2081"),
    ("Output Torque", "T\u2082 = T\u2081 \u00d7 i \u00d7 \u03b7   (\u03b7 \u2248 0.95 per spur stage)"),
    ("Lewis Bending Stress", "\u03c3 = F_t / (b \u00d7 m \u00d7 Y)   (tangential force, face width, module, form factor)"),
    ("PLA Strength (ACME)", "~52 MPa bulk; use 28 MPa for interlayer-critical features; SF \u2265 2.0"),
]
for cells_data in refs:
    add_table_row(tbl_ref, cells_data)

for row in tbl_ref.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(4.7)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════════
# "BE AN ENGINEER" CALLOUT
# ════════════��══════════════════════════════════��═══════════════════════════

# Create a single-cell table for the callout box
callout = doc.add_table(rows=1, cols=1)
callout.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = callout.rows[0].cells[0]
set_cell_shading(cell, "FFF8E1")

# Border the cell
tc_pr = cell._element.get_or_add_tcPr()
borders = tc_pr.makeelement(qn("w:tcBorders"), {})
for edge in ["top", "left", "bottom", "right"]:
    border = borders.makeelement(qn(f"w:{edge}"), {
        qn("w:val"): "single",
        qn("w:sz"): "8",
        qn("w:space"): "0",
        qn("w:color"): "D4A017",
    })
    borders.append(border)
tc_pr.append(borders)

p = cell.paragraphs[0]
run = p.add_run("Remember: Be an Engineer")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x8B, 0x6F, 0x00)

p = cell.add_paragraph()
run = p.add_run(
    "The AI is a tool you direct, not a teammate who does the work for you. "
    "The goal of this project is not a perfect PRD \u2014 it\u2019s demonstrating that "
    "you can manage the information available to AI systems and evaluate whether "
    "better context produces better engineering answers. If you can\u2019t explain why "
    "your context-augmented answer is better than the baseline, that\u2019s a problem "
    "regardless of how polished the output appears."
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x5D, 0x4A, 0x00)


# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════���══════════════════════════════════════════════════

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
